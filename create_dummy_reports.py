import docx
import os

# Create directory if it doesn't exist
dir_path = "relazioni_word/2024"
os.makedirs(dir_path, exist_ok=True)

# --- Dummy Document 1: Hydraulic Pump ---
doc1 = docx.Document()
doc1.add_paragraph("Relazione Intervento Pompa Idraulica")
doc1.add_paragraph("In data odierna, si è intervenuti sulla pompa idraulica modello H-500 a seguito di una segnalazione di bassa pressione.")
doc1.add_paragraph("All'arrivo sull'impianto, è stata riscontrata una perdita vistosa dal corpo della guarnizione principale.")
doc1.add_paragraph("Si è proceduto con lo smontaggio della pompa, la sostituzione della guarnizione e il rimontaggio completo. La pressione è tornata ai valori nominali.")
doc1.save(os.path.join(dir_path, "relazione_pompa.docx"))

# --- Dummy Document 2: Electrical Valve ---
doc2 = docx.Document()
doc2.add_paragraph("Manutenzione Valvola Elettrica V-201")
doc2.add_paragraph("Intervento di manutenzione programmata sulla valvola elettrica V-201.")
doc2.add_paragraph("L'attuatore è stato testato e i contatti elettrici sono stati puliti e serrati. La valvola risponde correttamente ai comandi del PLC.")
doc2.add_paragraph("Non sono state riscontrate anomalie durante l'intervento. La manutenzione è da considerarsi completata con successo.")
doc2.save(os.path.join(dir_path, "relazione_valvola.docx"))

# --- Dummy Document 3: Generic Anomaly ---
doc3 = docx.Document()
doc3.add_paragraph("Report Anomalie Generiche")
doc3.add_paragraph("Durante il consueto giro di ispezione, è stato notato un rumore anomalo proveniente dal riduttore del nastro trasportatore NT-03.")
doc3.add_paragraph("Il livello dell'olio è risultato basso. Si è provveduto al rabbocco con olio specifico come da manuale.")
doc3.add_paragraph("Il rumore è diminuito ma si consiglia un monitoraggio nei prossimi giorni.")
doc3.save(os.path.join(dir_path, "relazione_anomalia.docx"))

print("Dummy DOCX files created successfully in relazioni_word/2024/")
